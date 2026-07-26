"""Generate evaluation figures and Chapters 4–8 as a Word document."""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

import joblib
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from sklearn.metrics import (
  accuracy_score,
  confusion_matrix,
  f1_score,
  precision_score,
  recall_score,
  roc_auc_score,
  roc_curve,
)

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT / "src"))

from dataset_utils import (
  FEATURE_COLUMNS,
  build_balanced_dataset_from_parts,
  list_complete_parts,
  prepare_train_test_splits,
  ThresholdClassifier,
)

RESULTS_DIR = PROJECT_ROOT / "outputs" / "model_results"
FIGURES_DIR = RESULTS_DIR / "figures"
DOCS_DIR = PROJECT_ROOT / "docs"
MODELS_DIR = PROJECT_ROOT / "models"
CHAPTERS_MD = DOCS_DIR / "chapters-4-8.md"
OUTPUT_DOCX = DOCS_DIR / "Chapters_4_to_8.docx"

MODEL_NAMES = [
  "decision_tree",
  "logistic_regression",
  "xgboost",
  "random_forest",
]
DISPLAY_NAMES = {
  "decision_tree": "Decision Tree",
  "logistic_regression": "Logistic Regression",
  "xgboost": "XGBoost",
  "random_forest": "Random Forest",
}


def ensure_dirs() -> None:
  FIGURES_DIR.mkdir(parents=True, exist_ok=True)
  DOCS_DIR.mkdir(parents=True, exist_ok=True)


def load_test_data() -> tuple[pd.DataFrame, pd.Series]:
  parts = list_complete_parts()[:1]
  df, _, _registry = build_balanced_dataset_from_parts(parts=parts)
  _X_train, _y_train, X_test, y_test, _Xs_train, _Xs_test, _scaler = prepare_train_test_splits(df)
  return X_test, y_test


def load_tuned_model(name: str):
  path = MODELS_DIR / f"{name}_tuned.joblib"
  if not path.exists():
    raise FileNotFoundError(path)
  model = joblib.load(path)
  meta_path = MODELS_DIR / f"{name}_tuned_meta.json"
  if meta_path.exists():
    threshold = json.loads(meta_path.read_text(encoding="utf-8")).get("threshold", 0.5)
    if not isinstance(model, ThresholdClassifier):
      model = ThresholdClassifier(model, threshold=threshold)
  return model


def evaluate_all_models() -> list[dict]:
  X_test, y_test = load_test_data()
  results = []
  for name in MODEL_NAMES:
    model = load_tuned_model(name)
    y_pred = model.predict(X_test)
    y_proba = model.predict_proba(X_test)[:, 1]
    cm = confusion_matrix(y_test, y_pred).tolist()
    results.append(
      {
        "model": name,
        "display": DISPLAY_NAMES[name],
        "accuracy": accuracy_score(y_test, y_pred),
        "precision": precision_score(y_test, y_pred, zero_division=0),
        "recall": recall_score(y_test, y_pred, zero_division=0),
        "f1_score": f1_score(y_test, y_pred, zero_division=0),
        "roc_auc": roc_auc_score(y_test, y_proba),
        "confusion_matrix": cm,
        "y_pred": y_pred,
        "y_proba": y_proba,
      }
    )
  return results, y_test


def plot_confusion_matrix(cm: list, title: str, output_path: Path) -> None:
  plt.figure(figsize=(5.5, 4.5))
  sns.heatmap(
    cm,
    annot=True,
    fmt="d",
    cmap="Blues",
    cbar=True,
    xticklabels=["Legitimate", "Fraud"],
    yticklabels=["Legitimate", "Fraud"],
    annot_kws={"size": 12},
  )
  plt.xlabel("Predicted Label", fontsize=11)
  plt.ylabel("Actual Label", fontsize=11)
  plt.title(title, fontsize=12, fontweight="bold")
  plt.tight_layout()
  plt.savefig(output_path, dpi=200, bbox_inches="tight")
  plt.close()


def plot_roc_curves(results: list[dict], y_test: pd.Series, output_path: Path) -> None:
  plt.figure(figsize=(8, 6))
  for result in results:
    fpr, tpr, _ = roc_curve(y_test, result["y_proba"])
    plt.plot(
      fpr,
      tpr,
      linewidth=2,
      label=f"{result['display']} (AUC={result['roc_auc']:.3f})",
    )
  plt.plot([0, 1], [0, 1], "k--", linewidth=1, label="Random Classifier")
  plt.xlabel("False Positive Rate", fontsize=11)
  plt.ylabel("True Positive Rate", fontsize=11)
  plt.title("ROC Curves — Model Comparison", fontsize=12, fontweight="bold")
  plt.legend(loc="lower right", fontsize=9)
  plt.grid(alpha=0.3)
  plt.tight_layout()
  plt.savefig(output_path, dpi=200, bbox_inches="tight")
  plt.close()


def plot_metrics_comparison(metrics_df: pd.DataFrame, output_path: Path) -> None:
  plot_df = metrics_df.melt(
    id_vars="model",
    value_vars=["accuracy", "precision", "recall", "f1_score", "roc_auc"],
    var_name="metric",
    value_name="score",
  )
  plot_df["model"] = plot_df["model"].map(DISPLAY_NAMES)
  plt.figure(figsize=(10, 5.5))
  sns.barplot(data=plot_df, x="metric", y="score", hue="model", palette="Set2")
  plt.ylim(0, 1.08)
  plt.title("Comparative Model Performance Metrics", fontsize=12, fontweight="bold")
  plt.xlabel("Metric", fontsize=11)
  plt.ylabel("Score", fontsize=11)
  plt.legend(title="Model", bbox_to_anchor=(1.02, 1), loc="upper left")
  plt.tight_layout()
  plt.savefig(output_path, dpi=200, bbox_inches="tight")
  plt.close()


def plot_system_architecture(output_path: Path) -> None:
  fig, ax = plt.subplots(figsize=(12, 8))
  ax.set_xlim(0, 12)
  ax.set_ylim(0, 10)
  ax.axis("off")

  colors = {
    "presentation": "#1e3a5f",
    "inference": "#2563eb",
    "ml": "#059669",
    "data": "#7c3aed",
    "text": "white",
    "subtext": "#1f2937",
  }

  def box(x, y, w, h, title, items, color, fontsize=9):
    patch = FancyBboxPatch(
      (x, y),
      w,
      h,
      boxstyle="round,pad=0.02,rounding_size=0.08",
      linewidth=1.5,
      edgecolor=color,
      facecolor=color,
      alpha=0.92,
    )
    ax.add_patch(patch)
    ax.text(x + w / 2, y + h - 0.35, title, ha="center", va="top", fontsize=11, fontweight="bold", color=colors["text"])
    for i, item in enumerate(items):
      ax.text(x + 0.2, y + h - 0.75 - i * 0.38, f"• {item}", ha="left", va="top", fontsize=fontsize, color=colors["text"])

  box(
    0.5, 7.0, 11, 2.2,
    "Presentation Layer (Flask Web Application)",
    ["Overview Dashboard  (/index)", "Transaction Analysis UI  (/predict)", "REST API  (/api/predict, /api/metrics)", "Health Check  (/health)"],
    colors["presentation"],
  )
  box(
    0.5, 4.3, 11, 2.2,
    "Inference Pipeline (FraudDetectionPipeline)",
    ["Input Validation", "Feature Engineering", "Account Legitimacy Lookup (SQLite Registry)", "Probability Scoring & Risk Band Assignment"],
    colors["inference"],
  )
  box(
    0.5, 2.0, 11, 1.8,
    "Machine Learning Layer",
    ["Decision Tree Classifier (deployed)", "ThresholdClassifier wrapper", "StandardScaler (where applicable)"],
    colors["ml"],
  )
  box(
    0.5, 0.2, 11, 1.5,
    "Data Layer",
    ["Cifer-Fraud-Detection-Dataset-AF", "Cleaning & Feature Engineering", "50/50 Balanced Training Split", "Account Legitimacy Registry"],
    colors["data"],
  )

  for y_from, y_to in [(7.0, 6.5), (4.3, 4.0), (2.0, 1.7)]:
    arrow = FancyArrowPatch((6, y_from), (6, y_to), arrowstyle="-|>", mutation_scale=14, linewidth=1.8, color="#374151")
    ax.add_patch(arrow)

  ax.text(
    6, 9.55,
    "Figure 4.1: Three-Tier System Architecture — Fraud Detection Prototype",
    ha="center", va="center", fontsize=13, fontweight="bold", color=colors["subtext"],
  )
  ax.text(
    6, 9.15,
    "Offline training (Python scripts) → Deploy artefacts (joblib + SQLite) → Online inference (Gunicorn / Render)",
    ha="center", va="center", fontsize=9, color="#4b5563",
  )

  plt.tight_layout()
  plt.savefig(output_path, dpi=200, bbox_inches="tight", facecolor="white")
  plt.close()


def generate_figures() -> dict[str, Path]:
  print("Evaluating models and generating figures...", flush=True)
  results, y_test = evaluate_all_models()

  figure_map: dict[str, Path] = {}

  arch_path = FIGURES_DIR / "figure_4_1_system_architecture.png"
  plot_system_architecture(arch_path)
  figure_map["figure_4_1_system_architecture.png"] = arch_path
  print(f"  Created {arch_path.name}", flush=True)

  for result in results:
    fname = f"confusion_matrix_{result['model']}.png"
    out = FIGURES_DIR / fname
    plot_confusion_matrix(
      result["confusion_matrix"],
      f"Confusion Matrix — {result['display']}",
      out,
    )
    figure_map[fname] = out
    print(f"  Created {fname}", flush=True)

  roc_path = FIGURES_DIR / "roc_curves_comparison.png"
  plot_roc_curves(results, y_test, roc_path)
  figure_map["roc_curves_comparison.png"] = roc_path
  print(f"  Created {roc_path.name}", flush=True)

  metrics_rows = [
    {
      "model": r["model"],
      "accuracy": round(r["accuracy"], 4),
      "precision": round(r["precision"], 4),
      "recall": round(r["recall"], 4),
      "f1_score": round(r["f1_score"], 4),
      "roc_auc": round(r["roc_auc"], 4),
    }
    for r in results
  ]
  metrics_df = pd.DataFrame(metrics_rows).sort_values("accuracy", ascending=False)
  metrics_path = FIGURES_DIR / "metrics_comparison.png"
  plot_metrics_comparison(metrics_df, metrics_path)
  figure_map["metrics_comparison.png"] = metrics_path
  print(f"  Created {metrics_path.name}", flush=True)

  report = {
    "figures": {k: str(v) for k, v in figure_map.items()},
    "metrics": metrics_rows,
  }
  (FIGURES_DIR / "figure_manifest.json").write_text(json.dumps(report, indent=2), encoding="utf-8")
  return figure_map


def add_formatted_paragraph(doc: Document, text: str, style: str = "Normal") -> None:
  text = text.strip()
  if not text or text.startswith("*Note:"):
    return
  if text.startswith("# "):
    doc.add_heading(text[2:].strip(), level=1)
    return
  if text.startswith("## "):
    doc.add_heading(text[3:].strip(), level=2)
    return

  para = doc.add_paragraph(style=style)
  parts = re.split(r"(\*\*[^*]+\*\*)", text)
  for part in parts:
    if part.startswith("**") and part.endswith("**"):
      run = para.add_run(part[2:-2])
      run.bold = True
    else:
      part = re.sub(r"`([^`]+)`", r"\1", part)
      if part:
        para.add_run(part)


def parse_table_lines(lines: list[str]) -> tuple[list[str], list[list[str]]] | None:
  if len(lines) < 2 or not lines[0].strip().startswith("|"):
    return None
  header = [c.strip() for c in lines[0].strip("|").split("|")]
  rows = []
  for line in lines[2:]:
    if not line.strip().startswith("|"):
      break
    rows.append([c.strip() for c in line.strip("|").split("|")])
  return header, rows


def add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
  table = doc.add_table(rows=1, cols=len(header))
  table.style = "Table Grid"
  hdr_cells = table.rows[0].cells
  for i, col in enumerate(header):
    hdr_cells[i].text = col
    for paragraph in hdr_cells[i].paragraphs:
      for run in paragraph.runs:
        run.bold = True
  for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
      cells[i].text = val
  doc.add_paragraph()


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 5.5) -> None:
  if not image_path.exists():
    doc.add_paragraph(f"[Figure missing: {caption}]")
    return
  doc.add_paragraph()
  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  run = p.add_run()
  run.add_picture(str(image_path), width=Inches(width))
  cap = doc.add_paragraph(caption)
  cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
  for run in cap.runs:
    run.italic = True
    run.font.size = Pt(10)
  doc.add_paragraph()


def build_word_document(figure_map: dict[str, Path]) -> Path:
  print("Building Word document...", flush=True)
  text = CHAPTERS_MD.read_text(encoding="utf-8")
  doc = Document()

  style = doc.styles["Normal"]
  style.font.name = "Times New Roman"
  style.font.size = Pt(12)

  lines = text.splitlines()
  i = 0
  figure_insertions = {
    "Figure 4.1": figure_map.get("figure_4_1_system_architecture.png"),
    "Table 4.1": "table",
    "Figure 4.2": figure_map.get("metrics_comparison.png"),
    "Figure 4.3": figure_map.get("roc_curves_comparison.png"),
  }

  confusion_figures = [
    (name, figure_map.get(f"confusion_matrix_{name}.png"))
    for name in MODEL_NAMES
  ]

  while i < len(lines):
    line = lines[i]

    if line.strip().startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
      table_lines = []
      j = i
      while j < len(lines) and lines[j].strip().startswith("|"):
        table_lines.append(lines[j])
        j += 1
      parsed = parse_table_lines(table_lines)
      if parsed:
        header, rows = parsed
        if header and header[0].startswith("---"):
          i = j
          continue
        doc.add_paragraph("Table 4.1: Comparative model performance on the held-out test set")
        add_table(doc, header, rows)
        i = j
        continue

    if line.strip().startswith("```"):
      j = i + 1
      while j < len(lines) and not lines[j].strip().startswith("```"):
        j += 1
      i = j + 1
      continue

    if "**Figure 4.1:" in line:
      add_formatted_paragraph(doc, line.replace("**", ""))
      add_figure(doc, figure_map["figure_4_1_system_architecture.png"], "Figure 4.1: System architecture of the fraud detection prototype", width=6.2)
      i += 1
      continue

    if "**Figure 4.2:" in line:
      add_formatted_paragraph(doc, line.replace("**", ""))
      add_figure(
        doc,
        figure_map["metrics_comparison.png"],
        "Figure 4.2: Comparative model performance metrics (accuracy, precision, recall, F1-score, ROC-AUC)",
        width=6.2,
      )
      i += 1
      continue

    if "**Figure 4.3:" in line:
      add_formatted_paragraph(doc, line.replace("**", ""))
      add_figure(
        doc,
        figure_map["roc_curves_comparison.png"],
        "Figure 4.3: ROC curves for all trained classifiers on the held-out test set",
        width=6.0,
      )
      i += 1
      continue

    if line.strip().startswith("*[Insert"):
      i += 1
      continue

    if "## 4.6 Testing and Technical Validation" in line:
      add_formatted_paragraph(doc, line)
      i += 1
      doc.add_paragraph(
        "Figure 4.4 to Figure 4.7 present confusion matrices for each trained classifier on the held-out test set."
      )
      fig_num = 4
      for model_key, fig_path in confusion_figures:
        if fig_path and fig_path.exists():
          add_figure(
            doc,
            fig_path,
            f"Figure 4.{fig_num}: Confusion matrix — {DISPLAY_NAMES[model_key]}",
            width=4.2,
          )
          fig_num += 1
      continue

    if line.strip() == "---":
      i += 1
      continue

    if line.strip():
      add_formatted_paragraph(doc, line)
    i += 1

  doc.save(OUTPUT_DOCX)
  print(f"Saved {OUTPUT_DOCX}", flush=True)
  return OUTPUT_DOCX


def main() -> None:
  ensure_dirs()
  figure_map = generate_figures()
  build_word_document(figure_map)


if __name__ == "__main__":
  main()
